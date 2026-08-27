"""Export utilities for Cover Letter generation (DOCX, PDF, LaTeX)."""

from __future__ import annotations

import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF
from fpdf.enums import XPos, YPos

from wecanfindintern.career_ai.models import UserProfile


def export_docx(body_text: str, user_info: UserProfile) -> bytes:
    """Create a styled Word (.docx) document."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header: Candidate Name
    if user_info.full_name:
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_name = p_name.add_run(user_info.full_name)
        r_name.bold = True
        r_name.font.size = Pt(16)
        r_name.font.name = "Calibri"
        r_name.font.color.rgb = RGBColor(23, 64, 47)  # WeCanFindIntern dark-mint

        # Contact Info Line
        contact_parts = [
            user_info.email,
            user_info.phone,
            user_info.linkedin,
            user_info.address,
        ]
        contact_str = "  ·  ".join([p for p in contact_parts if p])
        if contact_str:
            p_contact = doc.add_paragraph()
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_contact = p_contact.add_run(contact_str)
            r_contact.font.size = Pt(10)
            r_contact.font.name = "Calibri"
            r_contact.font.color.rgb = RGBColor(100, 116, 105)
            p_contact.paragraph_format.space_after = Pt(18)

    # Letter Body Paragraphs
    for paragraph_text in body_text.split("\n\n"):
        clean_para = paragraph_text.strip()
        if not clean_para:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(clean_para)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(30, 41, 59)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


class StyledPDF(FPDF):
    def __init__(self, user_info: UserProfile):
        super().__init__()
        self.user_info = user_info
        self.set_auto_page_break(auto=True, margin=20)
        self.add_page()

    def header(self):
        if self.user_info.full_name:
            self.set_font("Helvetica", "B", 15)
            self.set_text_color(23, 64, 47)
            self.cell(0, 8, self.user_info.full_name, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")

            contact_parts = [
                self.user_info.email,
                self.user_info.phone,
                self.user_info.linkedin,
                self.user_info.address,
            ]
            contact_str = " | ".join([p for p in contact_parts if p])
            if contact_str:
                self.set_font("Helvetica", "", 9)
                self.set_text_color(100, 116, 105)
                self.cell(0, 5, contact_str, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
                self.ln(6)


def export_pdf(body_text: str, user_info: UserProfile) -> bytes:
    """Create a styled PDF document."""
    pdf = StyledPDF(user_info)
    pdf.set_font("Helvetica", "", 10.5)
    pdf.set_text_color(30, 41, 59)

    for para in body_text.split("\n\n"):
        clean_p = para.strip().replace("\u2013", "-").replace("\u2014", "--").replace("\u2019", "'")
        if not clean_p:
            continue
        pdf.multi_cell(0, 6, clean_p)
        pdf.ln(3)

    return bytes(pdf.output())


def export_latex(body_text: str, user_info: UserProfile, date_str: str = "[Date]") -> str:
    """Create a LaTeX source code (.tex) document."""
    # Sanitize LaTeX special characters
    def sanitize(text: str) -> str:
        replacements = {
            "\\": "\\textbackslash{}",
            "&": "\\&",
            "%": "\\%",
            "$": "\\$",
            "#": "\\#",
            "_": "\\_",
            "{": "\\{",
            "}": "\\}",
            "~": "\\textasciitilde{}",
            "^": "\\textasciicircum{}",
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text

    name = sanitize(user_info.full_name or "Applicant")
    email = sanitize(user_info.email)
    phone = sanitize(user_info.phone)
    linkedin = sanitize(user_info.linkedin)
    address = sanitize(user_info.address)

    body_escaped = "\n\n".join(
        [f"{sanitize(p.strip())}" for p in body_text.split("\n\n") if p.strip()]
    )

    return f"""\\documentclass[11pt,a4paper]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[margin=1in]{{geometry}}
\\usepackage{{hyperref}}
\\usepackage{{parskip}}

\\hypersetup{{
    colorlinks=true,
    linkcolor=teal,
    urlcolor=teal
}}

\\begin{{document}}

\\begin{{center}}
    {{\\LARGE \\textbf{{{name}}}}} \\\\[4pt]
    {email} {f'$\\cdot$ {phone}' if phone else ''} {f'$\\cdot$ {linkedin}' if linkedin else ''} {f'$\\cdot$ {address}' if address else ''}
\\end{{center}}

\\vspace{{1em}}
\\hrule
\\vspace{{1.5em}}

{body_escaped}

\\end{{document}}
"""
